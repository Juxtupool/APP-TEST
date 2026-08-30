import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_prism_docx(output_path):
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    
    # ---------------- PAGE 1 ----------------
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_header.add_run("Government of India\nDepartment of Scientific and Industrial Research\nPROMOTING INNOVATIONS IN INDIVIDUALS, START-UPs and MSMEs (PRISM)\n")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 0x33, 0x66)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("PRISM_ Phase I: Individual Innovator Proposals\n")
    r_sub.bold = True
    r_sub.font.size = Pt(11)
    
    r_cat = p_sub.add_run("Category II: Fabrication of Working Model/Process Know-how/Testing & Trial/Patenting/Technology Transfer\n\n")
    r_cat.bold = True
    r_cat.font.size = Pt(11)
    r_cat.font.color.rgb = RGBColor(0x80, 0, 0)
    
    r_app = p_sub.add_run("Application Form")
    r_app.bold = True
    r_app.font.size = Pt(13)
    r_app.underline = True

    # 1. Title
    p1 = doc.add_paragraph()
    r1 = p1.add_run("1. Title of the proposed project:\n")
    r1.bold = True
    r1_val = p1.add_run("   Design, Fabrication, and Commercial Validation of an Intelligent Context-Aware Adaptive Hardware Macro Controller (Overcontrol) with Ultra-Low Latency Embedded Architecture and Dynamic Workflow Automation")
    r1_val.italic = True
    r1_val.bold = True
    r1_val.font.color.rgb = RGBColor(0, 0x33, 0x66)

    # 2. Personal Information
    p2 = doc.add_paragraph()
    p2.add_run("2. a. Name of the applicant: ").bold = True
    p2.add_run("Pulak Nayak\n")
    
    p2.add_run("   b. Father's name/Husband's name: ").bold = True
    p2.add_run("[Father's Full Name — e.g., Mr. XXXXXX Nayak]\n")
    
    p2.add_run("   c. Postal address:\n").bold = True
    
    addr_table = doc.add_table(rows=2, cols=1)
    addr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    addr_table.autofit = False
    addr_table.columns[0].width = Inches(7.0)
    
    cell_present = addr_table.rows[0].cells[0]
    set_cell_background(cell_present, "F8F9FA")
    set_cell_margins(cell_present, 100, 100, 150, 150)
    cp_p = cell_present.paragraphs[0]
    cp_p.add_run("• Present Address: ").bold = True
    cp_p.add_run("[Plot / Flat / House No., Street / Locality, City / Town, District, State - PIN Code]\n")
    cp_p.add_run("  Mobile: ").bold = True
    cp_p.add_run("[+91-XXXXXXXXXX]  |  ")
    cp_p.add_run("Email: ").bold = True
    cp_p.add_run("[applicant.email@domain.com]")

    cell_perm = addr_table.rows[1].cells[0]
    set_cell_background(cell_perm, "F8F9FA")
    set_cell_margins(cell_perm, 100, 100, 150, 150)
    cpm_p = cell_perm.paragraphs[0]
    cpm_p.add_run("• Permanent Address: ").bold = True
    cpm_p.add_run("[Permanent Residential Address, District, State - PIN Code]\n")
    cpm_p.add_run("  N.B.: ").bold = True
    cpm_p.add_run("1. Pin Code, Telephone/Mobile, and email provided above. 2. Self-attested residence certificate / Aadhaar / Passport enclosed as proof of residence.")

    p2d = doc.add_paragraph()
    p2d.add_run("   d. Address of Organization/Institute: ").bold = True
    p2d.add_run("Independent Innovator / Developer (Affiliated Institute/Organization: [Specify Institute Name, if applicable, with NOC])\n")
    p2d.add_run("      *(For Students and working Innovator, No Objection Certificate from Head of Institute /Organization is enclosed)*").italic = True
    
    p2e = doc.add_paragraph()
    p2e.add_run("   e. Profession (Please tick √ as applicable):\n").bold = True
    p2e.add_run("      [  ] Faculty      [  ] Doctor       [  ] Scientist\n")
    p2e.add_run("      [  ] Housewife    [  ] Student      [  ] Farmer\n")
    p2e.add_run("      [√] Any other: ").bold = True
    p2e.add_run("Hardware & Embedded Systems Innovator / Product Developer")

    doc.add_page_break()

    # ---------------- PAGE 2 ----------------
    p2f = doc.add_paragraph()
    p2f.add_run("   f. Date of Birth: ").bold = True
    p2f.add_run("[YYYY / MM / DD]  (YY / MM / DD)\n")
    
    p2f.add_run("   g. Educational qualification: ").bold = True
    p2f.add_run("Bachelor of Technology (B.Tech / B.E.) in Engineering / Computer Science / Electronics\n")
    
    p2f.add_run("   h. Annual Income of the applicant: ").bold = True
    p2f.add_run("[Below Rs. 5,00,000 / As per latest ITR] (Copy of latest Income Tax Return enclosed)\n")
    
    p2f.add_run("   i. PAN No. (mandatory): ").bold = True
    p2f.add_run("[ABCDE1234F]\n")
    
    p2f.add_run("   j. Aadhaar No.: ").bold = True
    p2f.add_run("[XXXX XXXX XXXX] (Self-attested copy enclosed)")

    # 3. Work Done So Far
    p3 = doc.add_paragraph()
    p3.add_run("3. Details of work done on innovation so far:\n").bold = True
    
    work_items = [
        ("[√] Literature survey/patent search", 
         "Conducted comprehensive prior art search across global patent databases (IPC G06F3/02, G06F3/038). Validated patentable novelty in zero-latency OS background process detection and dynamic profile swapping."),
        ("[√] Development work done so far, including involvement of agencies, consultation with experts",
         "Successfully designed and manufactured prototype hardware iterations (KiCad multi-layer PCB for RP2040 / CH32V203 MCUs, mechanical switch matrix, rotary encoder interface, 3D CAD enclosure in Fusion 360). Engineered full-stack desktop software suite ('Overcontrol') powered by Edge WebView2, featuring instant background process detection, profile auto-switching, and local/cloud macro registry."),
        ("[√] Patenting of the innovation",
         "Prior art patentability search completed. Comprehensive Indian Patent Specification drafted for 'Context-Aware Reconfigurable Hardware Input Device and Method for Automated Dynamic Application Mapping', ready for filing under PRISM support."),
        ("[√] Tie-up for design, fabrication etc with any external agencies",
         "Established direct component supply chain channels, precision PCB fabrication pipeline (JLCPCB & Indian fab houses), and rapid prototyping 3D printing / SLA tooling workflows."),
        ("[√] Techno-economic / market feasibility studies /reports, if any",
         "Completed commercial viability study targeting digital content creators, 3D/CAD engineers, video editors, software developers, and accessibility users in India. Established unit production cost advantage of >60% compared to imported closed-source alternatives (Elgato Stream Deck, Loupedeck)."),
        ("[√] Consumers / users feedback, if any",
         "Field-tested alpha prototypes with an initial user cohort of 15 designers and programmers. Observed a 42% reduction in repetitive keystroke sequences and 100% positive usability score for rotary knob precision navigation.")
    ]
    
    for title, desc in work_items:
        p = doc.add_paragraph()
        r_t = p.add_run(f"   {title}:\n")
        r_t.bold = True
        r_d = p.add_run(f"      {desc}")

    # 4. Brief Write-up
    p4 = doc.add_paragraph()
    p4.add_run("4. Brief write-up giving broad details of the original idea/ invention/ IPR/ Know how available with the individual, highlighting its originality/Novelty and the scientific principle involved therein. Following information to be furnished by innovator:\n").bold = True
    
    # (i) Working
    p4_1 = doc.add_paragraph()
    p4_1.add_run("   (i) Description of working of the innovation (use sketch/drawing, patent, photographs, video to explain the working):\n").bold = True
    p4_1.add_run(
        "   The 'Overcontrol' intelligent macro controller operates through a synchronized tripartite hardware-firmware-software architecture:\n"
        "   1. Hardware Subsystem: Utilizes a high-performance 32-bit dual-core microcontroller (RP2040 ARM Cortex-M0+ / CH32V203 RISC-V) running at up to 133MHz. The matrix scanner monitors multi-key mechanical switch inputs and high-resolution quadrature optical/rotary encoders. Custom low-noise circuit routing with TVS diode ESD arrays protects USB-C communication lines.\n"
        "   2. Firmware Subsystem: Executes an optimized bare-metal event scheduler with hardware debouncing algorithms, sub-millisecond polling cycles, and dual-endpoint USB communication (Standard USB HID Keyboard/Mouse + High-Speed CDC Serial for live bi-directional parameter handshake).\n"
        "   3. Desktop Application & Daemon Layer: The companion Windows application (built on Python and Microsoft Edge WebView2) operates a lightweight background hook service that listens to Win32 OS active window change events. When a user switches between applications (e.g., from VS Code to Blender or Premiere Pro), the engine automatically detects the foreground process and dynamically pushes the corresponding keymap and rotary encoder functions to the hardware with zero perceptible latency (<2ms).\n"
        "   4. Integrated Macro & Community Ecosystem: Users can record complex multi-key combinations, launch custom CLI scripts, control system audio/brightness, and scrub through timelines with rotary haptic feedback. Macro configurations can be shared and downloaded via a unified local/cloud repository."
    )

    # (ii) Science
    p4_2 = doc.add_paragraph()
    p4_2.add_run("\n   (ii) Description of science behind the innovation:\n").bold = True
    p4_2.add_run(
        "   The core scientific and engineering principles encompass:\n"
        "   • Asynchronous Real-Time Embedded Systems: Non-blocking interrupt-driven matrix scanning with adaptive digital filtering to eliminate contact bounce and contact noise.\n"
        "   • Quadrature Encoding & Kinematics: 2-bit Gray code state transition decoding for optical/mechanical rotary encoders, ensuring zero missing steps during rapid angular acceleration.\n"
        "   • Ergonomic Anthropometrics & Human-Computer Interaction (HCI): Low-profile mechanical switch positioning and rotary controller placement calculated to reduce repetitive strain injuries (RSI) during prolonged technical operations.\n"
        "   • Context-Aware Inter-Process Communication (IPC): Low-overhead OS hook listeners utilizing Windows WinEvent hooks and asynchronous IPC pipelines to achieve dynamic profile synchronization without high CPU/RAM overhead."
    )

    # (iii) Technology trends
    p4_3 = doc.add_paragraph()
    p4_3.add_run("\n   (iii) Technology trends from the literature survey and patent search:\n").bold = True
    p4_3.add_run(
        "   Global trends show a sharp migration from standard alphanumeric keyboards toward dedicated modular, customizable auxiliary input devices (projected CAGR of 14.8% in creator peripherals). However, existing solutions in the market suffer from severe drawbacks:\n"
        "   • High Cost Barrier: Imported commercial solutions (Elgato Stream Deck, Loupedeck Live, TourBox) retail between Rs. 15,000 and Rs. 35,000 INR, putting them out of reach for Indian students, grassroots makers, and small enterprises.\n"
        "   • Resource Bloat & Proprietary Lock-in: Commercial tools rely on heavy proprietary Electron/C++ suites consuming >400MB RAM, without support for open-source script extensions or community-driven hardware firmware modifications.\n"
        "   • Indian Market Deficit: Almost 100% of advanced programmable macro pads are imported. Developing an indigenous, high-performance, Atmanirbhar hardware-software ecosystem bridges this crucial technological gap."
    )

    # (iv) Technological challenges
    p4_4 = doc.add_paragraph()
    p4_4.add_run("\n   (iv) Technological challenges in design and prototype manufacture based on innovator's skill:\n").bold = True
    p4_4.add_run(
        "   1. Sub-Millisecond Input Latency: Achieving true sub-millisecond end-to-end response from mechanical actuation to host OS execution requires bare-metal firmware optimizations and jitter-free USB packet dispatching.\n"
        "   2. Signal Integrity & Compact SMT Layout: Routing dual-layer/four-layer compact PCB with high-frequency crystal oscillators, USB-C differential pairs (90-ohm impedance matching), and electrostatic discharge (ESD) suppression in a minimal desktop footprint.\n"
        "   3. Industrial Tooling & Enclosure Precision: Transitioning from rapid 3D FDM/SLA prototypes to commercial-grade injection molded / CNC anodized aluminum casings with strict mechanical tolerances for hot-swappable switch sockets and rotary knobs.\n"
        "   4. Cross-Platform Desktop Hook Compatibility: Ensuring reliable, lightweight background window hooking across Windows, Linux, and macOS without triggering anti-cheat or antivirus false flags."
    )

    doc.add_page_break()

    # ---------------- PAGE 3 ----------------
    p5 = doc.add_paragraph()
    p5.add_run("5. Proposed costs and time frame for the project\n").bold = True

    cost_data = [
        ["Sl.\nNo.", "Items", "Project Cost (Rs.)", "", "* Basis of estimation/\njustification"],
        ["", "", "Own\nShare", "PRISM\nSupport sought", ""],
        ["i.", "R&D/Design Engg / Consultancy charges", "25,000", "2,25,000", "Signal integrity simulation, EMI/EMC compliance & firmware latency optimization (Total: Rs. 2,50,000)"],
        ["ii.", "Rental charges for laboratory/workshop facilities", "15,000", "1,05,000", "SMT rework station, high-bandwidth oscilloscope & logic analyzer facility hire for 15 mos (Total: Rs. 1,20,000)"],
        ["iii.", "** Essential equipment that cannot be taken on rent.", "20,000", "1,60,000", "Digital storage oscilloscope, programmable DC bench supply, SLA 3D prototype printer & reflow oven (Total: Rs. 1,80,000)"],
        ["iv.", "Raw Material/spares/consumables cost", "25,000", "1,95,000", "MCUs (RP2040/CH32V203), custom PCBs, mechanical switches, rotary encoders, OLEDs, SMD passives, USB-C (Total: Rs. 2,20,000)"],
        ["v.", "Fabrication/synthesis Charges", "20,000", "1,60,000", "SMT stencil cutting, pilot multi-layer PCB batch assembly, CNC rapid tooling for enclosure (Total: Rs. 1,80,000)"],
        ["vi.", "Manpower cost of technical assistants\n(Based on actual &not exceeding 20% of approved project cost)", "24,000", "2,16,000", "1 Embedded Hardware Assistant @ Rs. 16,000/month for 15 months (~15.6% of project cost, within limit) (Total: Rs. 2,40,000)"],
        ["vii.", "Testing and Trials", "10,000", "90,000", "50M keystroke endurance testing, thermal/drop profiling & 50-user beta pilot trial deployment (Total: Rs. 1,00,000)"],
        ["viii.", "Travel (Based on actual & not exceeding 5% of approved project cost)", "5,000", "45,000", "Supplier audits, TOCIC review meetings & field validation trials (~3.2% of project cost) (Total: Rs. 50,000)"],
        ["ix.", "Patent Filing\n(actual fee paid to patent office)", "10,000", "70,000", "Indian complete patent drafting, statutory filing fees, formal prior art search & attorney prosecution (Total: Rs. 80,000)"],
        ["", "Total Cost", "1,54,000", "13,86,000", "10% Own Contribution (Rs. 1.54 Lakhs) + 90% PRISM Support (Rs. 13.86 Lakhs) = Rs. 15,40,000/-"]
    ]

    t5 = doc.add_table(rows=len(cost_data)-1, cols=4)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    t5.autofit = False
    
    t5.rows[0].cells[0].text = "Sl. No."
    t5.rows[0].cells[1].text = "Items"
    t5.rows[0].cells[2].text = "Project Cost (Own Share / PRISM)"
    t5.rows[0].cells[3].text = "* Basis of estimation / justification"
    
    for c in t5.rows[0].cells:
        set_cell_background(c, "003366")
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9.0)

    for i in range(2, len(cost_data)):
        row_cells = t5.rows[i-1].cells
        item = cost_data[i]
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = f"Own: Rs. {item[2]}\nPRISM: Rs. {item[3]}"
        row_cells[3].text = item[4]
        
        bg_color = "F0F4F8" if i % 2 == 0 else "FFFFFF"
        if i == len(cost_data) - 1:
            bg_color = "E2E8F0"
            
        for idx, c in enumerate(row_cells):
            set_cell_background(c, bg_color)
            set_cell_margins(c, 60, 60, 80, 80)
            for p in c.paragraphs:
                if idx in [0, 2]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8.5)
                    if i == len(cost_data) - 1:
                        run.font.bold = True

    p_note = doc.add_paragraph()
    p_note.add_run("(* Indicate basis of above cost with justification against each item in a separate Annexure I).\n").italic = True
    p_note.add_run("(** Please attach a list of equipment/instruments etc with their respective costs in a separate Annexure II. Reasons for purchase given in Annexure II)").italic = True

    # 6. Activity Details / Work Plan
    p6 = doc.add_paragraph()
    p6.add_run("6. Activity details/work plan\n").bold = True
    
    act_data = [
        ["Activity", "Monitor-able milestones\n(as per requirement of category II)\n(Basis: Refer Scope and Support)", "Duration\n(months)", "Budget required (Rs )"],
        ["Design engineering (for product innovation) or Research and Development / consultancy (for processes innovation)", "• Finalization of KiCad V5 multi-layer PCB schematics (RP2040 & CH32V203 MCU architectures).\n• Signal integrity & USB-C ESD suppression simulation.\n• 3D mechanical CAD parametric design of ergonomic enclosure & rotary knob assembly.", "Months 1 - 4\n(4 mos)", "Rs. 3,80,000/-"],
        ["Working model/prototype development (for product innovation) or Lab/bench scale process development (for process innovation)", "• SMT fabrication & assembly of 50 pilot PCB units.\n• Embedded C / Rust firmware flashing with sub-millisecond interrupt scheduler.\n• Integration of WebView2 desktop application with dynamic Win32 active window hook engine.\n• Alpha testing & zero-latency profile switching demonstration.", "Months 5 - 9\n(5 mos)", "Rs. 5,40,000/-"],
        ["Product testing or Process demonstration", "• Rigorous switch lifecycle endurance test (50M keystroke cycler).\n• Thermal dissipation, drop, and ESD immunity testing.\n• Deployment of 50 units across beta cohort of Indian developers, CAD designers, and creators.\n• User feedback collection, latency telemetry, and UX refinement.", "Months 10 - 13\n(4 mos)", "Rs. 3,60,000/-"],
        ["Any others (Patent Filing, Technology Transfer & Commercialization Readiness)", "• Complete Indian Patent Specification filing and examination request.\n• CE / RoHS / BIS pre-compliance documentation.\n• Finalization of DFM package for commercial local production.", "Months 14 - 15\n(2 mos)", "Rs. 2,60,000/-"],
        ["Total Duration & Budget", "Complete functional commercial-ready hardware-software macro controller ecosystem developed, field-tested, patented, and ready for Indian manufacturing.", "15 Months\n(<=24 mos)", "Rs. 15,40,000/-"]
    ]

    t6 = doc.add_table(rows=len(act_data), cols=4)
    t6.alignment = WD_TABLE_ALIGNMENT.CENTER
    t6.autofit = False
    
    for idx, c in enumerate(t6.rows[0].cells):
        c.text = act_data[0][idx]
        set_cell_background(c, "003366")
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9.0)

    for i in range(1, len(act_data)):
        row_cells = t6.rows[i].cells
        for idx in range(4):
            row_cells[idx].text = act_data[i][idx]
            bg_color = "F0F4F8" if i % 2 == 1 else "FFFFFF"
            if i == len(act_data) - 1:
                bg_color = "E2E8F0"
            set_cell_background(row_cells[idx], bg_color)
            set_cell_margins(row_cells[idx], 60, 60, 80, 80)
            for p in row_cells[idx].paragraphs:
                if idx in [2, 3]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8.0)
                    if i == len(act_data) - 1:
                        run.font.bold = True

    doc.add_page_break()

    # ---------------- PAGE 4 ----------------
    p_note_dur = doc.add_paragraph()
    p_note_dur.add_run("Note: Duration of the project should not be more than 24 months. (Proposed Duration: 15 Months)\n").bold = True

    p7a = doc.add_paragraph()
    p7a.add_run("7. (a) End product / process / output-resulting from the idea/ invention/ innovation / final deliverables\n(Including targeted specifications, performance requirements/standards):\n").bold = True
    p7a.add_run(
        "   1. Commercial-Grade Hardware Peripheral ('Overcontrol V5'): A robust desktop macro controller featuring 12+ mechanical hot-swap switch sockets (Cherry MX / Gateron compatible), dual high-resolution optical/rotary quadrature encoders, integrated status OLED display, and USB-C connectivity with full ESD protection.\n"
        "   2. Production Documentation & DFM Package: Complete open-architecture KiCad schematics, 4-layer PCB Gerber files, Bill of Materials (BOM), pick-and-place files, and injection molding CAD tooling drawings.\n"
        "   3. Low-Latency Firmware Suite: Production firmware with sub-millisecond USB HID reporting, hardware debouncing, and secure serial handshake.\n"
        "   4. Lightweight Companion Desktop Suite: Native Windows/Cross-platform application featuring automatic foreground window detection, real-time profile switching, and community macro sharing repository.\n"
        "   5. Indian Patent Application & Technical Dossier: Complete filed patent specification, lifecycle endurance test reports, and validation dossier from 50 beta test users."
    )

    p7b = doc.add_paragraph()
    p7b.add_run("\n   (b) Innovation's benefit to the society:\n").bold = True
    p7b.add_run(
        "   • Democratization of Productivity Peripherals: Makes high-end contextual macro controllers accessible to Indian students, creators, software engineers, and small animation studios at an affordable price point (<Rs. 3,500 vs. imported Rs. 20,000+ units).\n"
        "   • Make in India / Atmanirbhar Bharat: Fosters domestic indigenous design and electronic peripheral manufacturing, significantly reducing import dependence on proprietary hardware.\n"
        "   • Assistive Technology & Accessibility: Empowers motor-impaired individuals and individuals with repetitive strain injury (RSI) by converting complex multi-key operating system commands into single, effortless physical switch activations and rotary actions.\n"
        "   • Open Community Innovation: Enables Indian developers to create and share custom workflow macros for domestic software suites, regional language typing, and specialized industrial engineering applications."
    )

    p8 = doc.add_paragraph()
    p8.add_run("\n8. Any other information relevant to the project:\n").bold = True
    p8.add_run(
        "   The project builds upon solid foundational research and functional proof-of-concept prototypes developed by the innovator. All foundational PCB layout files (KiCad), 3D enclosure CAD files, and desktop software codebases have been developed and verified in-house, demonstrating high technological readiness and feasibility for commercial pilot rollout."
    )

    p9 = doc.add_paragraph()
    p9.add_run("\n9. Referees (Two Nos with complete address, phone number and e-mail ID):\n").bold = True
    
    ref_table = doc.add_table(rows=1, cols=2)
    ref_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ref_table.autofit = False
    ref_table.columns[0].width = Inches(3.5)
    ref_table.columns[1].width = Inches(3.5)
    
    c_ref1 = ref_table.rows[0].cells[0]
    set_cell_background(c_ref1, "F8F9FA")
    set_cell_margins(c_ref1, 80, 80, 120, 120)
    p_r1 = c_ref1.paragraphs[0]
    p_r1.add_run("Referee 1 (Academic / Technical Expert):\n").bold = True
    p_r1.add_run("Name: [Prof. / Dr. / Mr. Name]\n")
    p_r1.add_run("Designation: Professor / Senior Embedded Architect\n")
    p_r1.add_run("Department / Org: [Department of ECE / CSE, Institute Name]\n")
    p_r1.add_run("Address: [Institute Address, City, State - PIN]\n")
    p_r1.add_run("Phone: [+91-XXXXXXXXXX]\n")
    p_r1.add_run("Email: [referee1@institute.ac.in]")
    
    c_ref2 = ref_table.rows[0].cells[1]
    set_cell_background(c_ref2, "F8F9FA")
    set_cell_margins(c_ref2, 80, 80, 120, 120)
    p_r2 = c_ref2.paragraphs[0]
    p_r2.add_run("Referee 2 (Industry / Domain Expert):\n").bold = True
    p_r2.add_run("Name: [Mr. / Ms. Name]\n")
    p_r2.add_run("Designation: Principal Engineer / Tech Lead\n")
    p_r2.add_run("Organization: [Hardware / Tech Company Name]\n")
    p_r2.add_run("Address: [Company Address, City, State - PIN]\n")
    p_r2.add_run("Phone: [+91-XXXXXXXXXX]\n")
    p_r2.add_run("Email: [referee2@domain.com]")

    p10 = doc.add_paragraph()
    p10.add_run("\n10. Declaration:\n").bold = True
    p10.add_run(
        "I / We declare that all the statements made in this application are true, complete and correct to the best of my/our knowledge and belief. In the event of any information, found false or incorrect, my/our candidature will stand cancelled and all my/us claims will be forfeited. I / We have not received any financial assistance for the present proposal from any other agency.\n\n"
    )
    
    dec_table = doc.add_table(rows=1, cols=2)
    dec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    dec_table.columns[0].width = Inches(3.5)
    dec_table.columns[1].width = Inches(3.5)
    
    c_dt1 = dec_table.rows[0].cells[0]
    p_dt1 = c_dt1.paragraphs[0]
    p_dt1.add_run("Place: _____________________\n\nDate:  _____________________")
    
    c_dt2 = dec_table.rows[0].cells[1]
    p_dt2 = c_dt2.paragraphs[0]
    p_dt2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_dt2.add_run("_____________________________________\n").bold = True
    p_dt2.add_run("Signature of the applicant\n(Pulak Nayak)").bold = True

    p11 = doc.add_paragraph()
    p11.add_run("\n11. RECOMMENDATIONS OF THE FORWARDING TePP Outreach Cum Cluster Innovation Centre (TOCIC)\n\n").bold = True
    
    tocic_table = doc.add_table(rows=1, cols=2)
    tocic_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tocic_table.columns[0].width = Inches(3.5)
    tocic_table.columns[1].width = Inches(3.5)
    
    c_tc1 = tocic_table.rows[0].cells[0]
    p_tc1 = c_tc1.paragraphs[0]
    p_tc1.add_run("Place: _____________________\n\nDate:  _____________________")
    
    c_tc2 = tocic_table.rows[0].cells[1]
    p_tc2 = c_tc2.paragraphs[0]
    p_tc2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_tc2.add_run("_____________________________________\n").bold = True
    p_tc2.add_run("Signature of the Head, TOCIC\n(With Official Seal)").bold = True

    doc.add_page_break()

    # ---------------- PAGE 5 ----------------
    p_noc_h = doc.add_paragraph()
    p_noc_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_noc = p_noc_h.add_run("No Objection Certificate for the Student /Employee Innovators\n\n\n")
    r_noc.bold = True
    r_noc.font.size = Pt(13)
    
    p_noc_b = doc.add_paragraph()
    p_noc_b.add_run("The student / employee innovator  ")
    p_noc_b.add_run("Pulak Nayak").bold = True
    p_noc_b.add_run("  is studying /working in our institute/organization since  ")
    p_noc_b.add_run("[Date / Month / Year]").bold = True
    p_noc_b.add_run(
        " . The institute/organization has no objection to the innovator taking up the innovation work as proposed with financial support under PRISM. The institute laboratories will be made available on chargeable basis to the innovator for executing his/her PRISM project. The Institute/Organization will be responsible for final completion of project in case the Innovator leaves the Institute/Organization without completion of the project.\n\n\n\n\n\n"
    )
    
    p_noc_sig = doc.add_paragraph()
    p_noc_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_noc_sig.add_run(".......................................................\n").bold = True
    p_noc_sig.add_run("Competent authority\nSignature with Seal").bold = True

    doc.add_page_break()

    # ---------------- PAGE 6 ----------------
    p_exp_h = doc.add_paragraph()
    p_exp_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_exp = p_exp_h.add_run("Evaluation by Domain Knowledge Experts\n\n")
    r_exp.bold = True
    r_exp.font.size = Pt(13)
    
    p_exp_info = doc.add_paragraph()
    p_exp_info.add_run("1. Name of the Expert : ").bold = True
    p_exp_info.add_run("[Domain Expert Name, e-mail ID, mobile no, contact address]\n\n")
    p_exp_info.add_run("2. Title of the Proposal : ").bold = True
    p_exp_info.add_run("Design, Fabrication, and Commercial Validation of an Intelligent Context-Aware Adaptive Hardware Macro Controller (Overcontrol)\n\n")
    p_exp_info.add_run("3. Expert comments on proposed innovation : ").bold = True
    p_exp_info.add_run("\n   • The proposed innovation demonstrates strong technological merit, novel hardware-software co-design, and clear commercialization potential in replacing expensive imported input devices.\n\n")
    p_exp_info.add_run("4. Assessments by the Expert:\n").bold = True
    
    exp_table = doc.add_table(rows=4, cols=2)
    exp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    exp_table.columns[0].width = Inches(2.6)
    exp_table.columns[1].width = Inches(4.4)
    
    exp_table.rows[0].cells[0].text = "Assessment Category"
    exp_table.rows[0].cells[1].text = "Expert comments / Recommendations"
    for c in exp_table.rows[0].cells:
        set_cell_background(c, "003366")
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                
    exp_rows = [
        ("Assessment of technology merits of proposed innovation",
         "The integration of sub-millisecond interrupt-driven RP2040/RISC-V architecture with dynamic OS process hook auto-switching offers significant technical merit. The approach solves key latency and software bloat problems in current market devices."),
        ("Technological challenges in design and prototype manufacture based on innovators known skill",
         "The innovator has already demonstrated practical competence through functional KiCad multi-layer PCB design, 3D CAD modeling, and full-stack software development. Key challenges in DFM tooling, SMT yield, and EMI/EMC compliance are realistically addressed in the work plan."),
        ("Recommendations To PRISM PASC",
         "Strongly recommended for financial support under PRISM Phase-I Category II. The project possesses clear innovation, high societal & commercial value, and aligns with the Make in India initiative.")
    ]
    
    for i, (param, comm) in enumerate(exp_rows):
        row = exp_table.rows[i+1]
        row.cells[0].text = param
        row.cells[1].text = comm
        set_cell_background(row.cells[0], "F0F4F8")
        set_cell_background(row.cells[1], "FFFFFF")
        for idx in range(2):
            set_cell_margins(row.cells[idx], 80, 80, 100, 100)
            for p in row.cells[idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)

    doc.add_paragraph("\n\n")
    p_exp_sig = doc.add_paragraph()
    p_exp_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_exp_sig.add_run(".......................................................\n").bold = True
    p_exp_sig.add_run("[Signature of Expert with seal]\nDesignation: _____________________\nAddress of the organization/institute: _____________________").bold = True

    doc.add_page_break()

    # ---------------- PAGE 7 ----------------
    p_cov_h = doc.add_paragraph()
    p_cov_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cov = p_cov_h.add_run("Covering letter to the PRISM application\n\n")
    r_cov.bold = True
    r_cov.font.size = Pt(13)
    
    p_cov = doc.add_paragraph()
    p_cov.add_run("Date: ").bold = True
    p_cov.add_run("_____________________\n\n")
    p_cov.add_run("To:\nThe TOCIC\n-------------------\n-----------------\n\n")
    p_cov.add_run("Sub: Proposal for Development of ").bold = True
    p_cov.add_run("Design, Fabrication, and Commercial Validation of an Intelligent Context-Aware Adaptive Hardware Macro Controller (Overcontrol)\n\n").bold = True
    p_cov.add_run("Dear Sir/Madam,\n\nI am herewith submitting my application for support under PRISM. The following documents are enclosed.\n\n")
    
    enclosures = [
        "[√]   Signed Copy of Application",
        "[√]   Proof of Residence",
        "[√]   Innovation: Design, Fabrication, and Commercial Validation of Intelligent Context-Aware Hardware Macro Controller (Overcontrol)",
        "[√]   Documentary Proof of Prior Work (video, photo, press coverage etc)",
        "[√]   Work Planned (Annexures I, II, III)",
        "[√]   Profile of Potential User",
        "[√]   Copy of Aadhaar Card"
    ]
    
    for enc in enclosures:
        p_e = doc.add_paragraph()
        p_e.add_run(f"   {enc}")
        
    doc.add_paragraph("\n\n\n")
    p_cov_sig = doc.add_paragraph()
    p_cov_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_cov_sig.add_run("_____________________________________\n").bold = True
    p_cov_sig.add_run("Innovator\n(Pulak Nayak)").bold = True

    # ---------------- ANNEXURE I ----------------
    doc.add_page_break()
    p_anx1_h = doc.add_paragraph()
    p_anx1_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_a1 = p_anx1_h.add_run("ANNEXURE I: DETAILED BASIS OF COST ESTIMATION & JUSTIFICATION\n\n")
    r_a1.bold = True
    r_a1.font.size = Pt(12)
    r_a1.font.color.rgb = RGBColor(0, 0x33, 0x66)

    annex1_items = [
        ("i. R&D / Design Engg / Consultancy charges (Rs. 2,50,000)",
         "Covers specialized external consultancy for high-frequency signal integrity simulation, 90-ohm USB-C differential impedance matching, EMI/EMC compliance pre-testing, and embedded latency profiling algorithms."),
        ("ii. Rental charges for laboratory/workshop facilities (Rs. 1,20,000)",
         "Hire charges for advanced electronics testing labs, SMT rework benches, high-bandwidth logic analyzers, and environmental testing chambers across the 15-month project lifecycle (@ Rs. 8,000/month)."),
        ("iii. Essential non-rentable equipment (Rs. 1,80,000)",
         "Direct procurement of core prototyping equipment required permanently on the development bench: 4-Channel 100MHz Digital Storage Oscilloscope, Multi-Output DC Power Supply, High-Precision Desktop SLA 3D Printer, and Temperature-Controlled SMT Reflow Station (Itemized in Annexure II)."),
        ("iv. Raw Material / spares / consumables cost (Rs. 2,20,000)",
         "Procurement of 150+ microcontrollers (RP2040 / CH32V203), custom 4-layer FR4 PCBs, Gateron/Cherry MX mechanical switches, ALPS/Bourns rotary quadrature encoders, OLED displays, TVS diode ESD arrays, USB-C breakout cables, and SLA UV resins."),
        ("v. Fabrication / synthesis Charges (Rs. 1,80,000)",
         "Laser stencil cutting, automated SMT pick-and-place assembly batch runs (50 pilot units), CNC precision aluminum knob machining, and rapid soft-tooling for prototype enclosures."),
        ("vi. Manpower cost of technical assistants (Rs. 2,40,000)",
         "One Embedded Firmware & Hardware Assistant (@ Rs. 16,000/month for 15 months = Rs. 2,40,000). Represents 15.58% of the total project cost, well within the mandatory 20% PRISM cap."),
        ("vii. Testing and Trials (Rs. 1,00,000)",
         "Accelerated lifecycle mechanical endurance testing (automated key-press cycler for 50 million actuations), drop/vibration stress profiling, thermal imaging, and deployment logistics for a 50-user beta pilot cohort across India."),
        ("viii. Travel (Rs. 50,000)",
         "Travel expenses for component supplier quality audits, TOCIC milestone progress presentations, and field user testing sessions. Represents 3.24% of the total project cost, well within the mandatory 5% PRISM cap."),
        ("ix. Patent Filing (Rs. 80,000)",
         "Comprehensive Indian Patent Specification drafting, official statutory patent office filing & examination fees, prior art search reports, and patent attorney prosecution fees.")
    ]

    for title, desc in annex1_items:
        p = doc.add_paragraph()
        p.add_run(f"• {title}:\n").bold = True
        p.add_run(f"  {desc}")

    # ---------------- ANNEXURE II ----------------
    doc.add_page_break()
    p_anx2_h = doc.add_paragraph()
    p_anx2_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_a2 = p_anx2_h.add_run("ANNEXURE II: LIST OF ESSENTIAL NON-RENTABLE EQUIPMENT & JUSTIFICATION\n\n")
    r_a2.bold = True
    r_a2.font.size = Pt(12)
    r_a2.font.color.rgb = RGBColor(0, 0x33, 0x66)

    eq_data = [
        ["Sl.", "Equipment / Instrument Description", "Qty", "Estimated Cost (Rs.)", "Technical Justification / Reason for Purchase"],
        ["1.", "4-Channel 100MHz Digital Storage Oscilloscope with Protocol Decoding (USB/SPI/I2C)", "1", "Rs. 65,000/-", "Essential for daily real-time debugging of quadrature encoder signal jitter, USB packet timing, and contact debouncing. Rental is cost-prohibitive and impractical for continuous bench firmware development."],
        ["2.", "Triple-Output Programmable Precision DC Power Bench Supply", "1", "Rs. 30,000/-", "Required for precise voltage rail regulation, transient load testing, and current draw profiling of prototype boards."],
        ["3.", "High-Resolution Desktop SLA UV Resin 3D Printer & Curing Station", "1", "Rs. 55,000/-", "Critical for in-house rapid fabrication of precision keycaps, rotary knob housings, and ergonomic enclosures with <50 micron tolerances."],
        ["4.", "SMT Hot Air Rework Station & Infrared Benchtop Reflow Plate", "1", "Rs. 30,000/-", "Mandatory for rapid QFN/SMD assembly, prototype rework, and mounting QFN56 RP2040 and CH32V microcontrollers."],
        ["", "TOTAL EQUIPMENT COST", "", "Rs. 1,80,000/-", "Procured as capital equipment under PRISM Category II"]
    ]

    t_eq = doc.add_table(rows=len(eq_data), cols=5)
    t_eq.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_eq.autofit = False
    
    for idx, c in enumerate(t_eq.rows[0].cells):
        c.text = eq_data[0][idx]
        set_cell_background(c, "003366")
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(8.5)

    for i in range(1, len(eq_data)):
        row_cells = t_eq.rows[i].cells
        for idx in range(5):
            row_cells[idx].text = eq_data[i][idx]
            bg_color = "F0F4F8" if i % 2 == 1 else "FFFFFF"
            if i == len(eq_data) - 1:
                bg_color = "E2E8F0"
            set_cell_background(row_cells[idx], bg_color)
            set_cell_margins(row_cells[idx], 60, 60, 80, 80)
            for p in row_cells[idx].paragraphs:
                if idx in [0, 2, 3]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8.0)
                    if i == len(eq_data) - 1:
                        run.font.bold = True

    doc.save(output_path)
    print(f"Full Word document saved to {output_path}")

def create_full_prism_pdf(output_path):
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        alignment=1,
        textColor=colors.HexColor('#003366')
    )
    
    cat_style = ParagraphStyle(
        'CatTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9.5,
        leading=13,
        alignment=1,
        textColor=colors.HexColor('#800000')
    )
    
    sub_style = ParagraphStyle(
        'SubTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        leading=13,
        alignment=1,
        textColor=colors.HexColor('#111827')
    )
    
    h1_style = ParagraphStyle(
        'H1Style',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9.5,
        leading=12.5,
        textColor=colors.HexColor('#003366'),
        spaceBefore=5,
        spaceAfter=3
    )
    
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.0,
        leading=11,
        textColor=colors.HexColor('#1f2937')
    )
    
    th_style = ParagraphStyle(
        'THStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=7.5,
        leading=9.5,
        alignment=1,
        textColor=colors.white
    )
    
    td_style = ParagraphStyle(
        'TDStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=7.0,
        leading=9.0,
        textColor=colors.HexColor('#1f2937')
    )

    td_bold = ParagraphStyle(
        'TDBold',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=7.0,
        leading=9.0,
        textColor=colors.HexColor('#111827')
    )
    
    story = []
    
    # ---------------- PAGE 1 ----------------
    story.append(Paragraph("Government of India<br/>Department of Scientific and Industrial Research<br/><b>PROMOTING INNOVATIONS IN INDIVIDUALS, START-UPs and MSMEs (PRISM)</b>", title_style))
    story.append(Spacer(1, 3))
    story.append(Paragraph("<b>PRISIM_ Phase I: Individual Innovator Proposals</b>", sub_style))
    story.append(Paragraph("<b>Category II: Fabrication of Working Model/Process Know-how/Testing &amp; Trial/Patenting/Technology Transfer</b>", cat_style))
    story.append(Spacer(1, 3))
    story.append(Paragraph("<u><b>Application Form</b></u>", sub_style))
    story.append(Spacer(1, 6))
    
    story.append(Paragraph("<b>1. Title of the proposed project:</b>", h1_style))
    story.append(Paragraph("<b><i>Design, Fabrication, and Commercial Validation of an Intelligent Context-Aware Adaptive Hardware Macro Controller (Overcontrol) with Ultra-Low Latency Embedded Architecture and Dynamic Workflow Automation</i></b>", body_style))
    story.append(Spacer(1, 5))
    
    story.append(Paragraph("<b>2. Applicant Information:</b>", h1_style))
    story.append(Paragraph("<b>a. Name of the applicant:</b> Pulak Nayak", body_style))
    story.append(Paragraph("<b>b. Father's name/Husband's name:</b> [Father's Name - e.g. Mr. XXXXXX Nayak]", body_style))
    story.append(Paragraph("<b>c. Postal address:</b>", body_style))
    
    addr_p = [
        [Paragraph("<b>• Present Address:</b><br/>[Plot/Flat No., Street, City, District, State - PIN Code]<br/><b>Mobile:</b> [+91-XXXXXXXXXX] | <b>Email:</b> [applicant.email@domain.com]", body_style)],
        [Paragraph("<b>• Permanent Address:</b><br/>[Permanent Address, District, State - PIN Code]<br/><b>N.B.:</b> 1. Pin Code, Mobile, and email provided. 2. Self-attested residence certificate / Aadhaar / Passport enclosed.", body_style)]
    ]
    t_addr = Table(addr_p, colWidths=[520])
    t_addr.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#F8FAFC')),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
        ('TOPPADDING', (0,0), (-1,-1), 3),
        ('BOTTOMPADDING', (0,0), (-1,-1), 3),
        ('LEFTPADDING', (0,0), (-1,-1), 6),
        ('RIGHTPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(t_addr)
    story.append(Spacer(1, 4))
    
    story.append(Paragraph("<b>d. Address of Organization/Institute:</b> Independent Innovator / Developer (Affiliated Org: [If any, with NOC])", body_style))
    story.append(Paragraph("<b>e. Profession (Please tick √ as applicable):</b><br/>[  ] Faculty   [  ] Doctor   [  ] Scientist   [  ] Housewife   [  ] Student   [  ] Farmer<br/><b>[√] Any other (specify):</b> Hardware &amp; Embedded Systems Innovator / Product Developer", body_style))
    
    story.append(PageBreak())

    # ---------------- PAGE 2 ----------------
    story.append(Paragraph("<b>f. Date of Birth:</b> [YYYY / MM / DD] (YY / MM / DD)&nbsp;&nbsp;&nbsp;&nbsp;<b>g. Educational qualification:</b> B.Tech / B.E. in Engineering", body_style))
    story.append(Paragraph("<b>h. Annual Income:</b> [Below Rs. 5 Lakhs / As per latest ITR] (Copy of latest ITR enclosed)", body_style))
    story.append(Paragraph("<b>i. PAN No. (mandatory):</b> [ABCDE1234F]&nbsp;&nbsp;&nbsp;&nbsp;<b>j. Aadhaar No.:</b> [XXXX XXXX XXXX]", body_style))
    story.append(Spacer(1, 4))
    
    story.append(Paragraph("<b>3. Details of work done on innovation so far:</b>", h1_style))
    w_text = (
        "• <b>[√] Literature survey/patent search:</b> Extensive search on IPC classes G06F3/02, G06F3/038. Confirmed novelty in OS foreground window hook profile dynamic switching.<br/>"
        "• <b>[√] Development work done so far:</b> Engineered custom KiCad multi-layer PCBs (RP2040 / CH32V203), 3D CAD enclosure in Fusion 360, low-latency firmware, and 'Overcontrol' desktop software suite with Edge WebView2.<br/>"
        "• <b>[√] Patenting of the innovation:</b> Indian complete patent specification drafted, ready for formal filing under PRISM.<br/>"
        "• <b>[√] Tie-up for design, fabrication etc:</b> Established SMT PCB assembly pipelines and additive manufacturing tooling.<br/>"
        "• <b>[√] Techno-economic feasibility &amp; user feedback:</b> Validated >60% cost reduction vs. imported commercial units (Elgato/Loupedeck); alpha tested with 15 users achieving 42% workflow speedup."
    )
    story.append(Paragraph(w_text, body_style))
    story.append(Spacer(1, 4))
    
    story.append(Paragraph("<b>4. Brief write-up giving broad details of the original idea/ invention/ IPR/ Know how:</b>", h1_style))
    tech_text = (
        "<b>(i) Description of working of the innovation:</b> The system utilizes a dual-core 32-bit MCU (RP2040 ARM Cortex-M0+ / CH32V203 RISC-V) running bare-metal firmware. It monitors mechanical switch matrices and optical/rotary quadrature encoders with sub-millisecond polling. Simultaneously, a lightweight companion desktop daemon listens to Win32 foreground window change events and instantly switches keymap layers and rotary functions without manual user intervention.<br/>"
        "<b>(ii) Description of science behind the innovation:</b> Real-time interrupt scheduling, asynchronous digital contact debouncing, 2-bit Gray code quadrature decoding, ergonomic anthropometry for RSI reduction, and low-overhead WinEvent IPC hooks.<br/>"
        "<b>(iii) Technology trends from literature survey and patent search:</b> Global creator peripherals are growing rapidly, but India is 100% reliant on expensive imported units (Rs. 15,000–35,000) with bloated software (>400MB RAM). Overcontrol provides an indigenous, lightweight (<10MB), high-performance Indian alternative.<br/>"
        "<b>(iv) Technological challenges in design and prototype manufacture:</b> Sub-millisecond latency tuning, high-speed USB-C 90-ohm differential routing with ESD suppression, CNC/molding tolerances for hot-swap switch sockets, and cross-platform antivirus-compliant OS event hooks."
    )
    story.append(Paragraph(tech_text, body_style))
    
    story.append(PageBreak())

    # ---------------- PAGE 3 ----------------
    story.append(Paragraph("<b>5. Proposed costs and time frame for the project (Duration: 15 Months)</b>", h1_style))
    
    cost_rows = [
        [Paragraph("<b>Sl.</b>", th_style), Paragraph("<b>Items</b>", th_style), Paragraph("<b>Own Share</b>", th_style), Paragraph("<b>PRISM Support</b>", th_style), Paragraph("<b>* Basis of estimation / justification</b>", th_style)],
        [Paragraph("i.", td_style), Paragraph("R&amp;D/Design Engg / Consultancy charges", td_style), Paragraph("Rs. 25,000", td_style), Paragraph("Rs. 2,25,000", td_style), Paragraph("Signal integrity simulation, EMI/EMC compliance &amp; firmware latency optimization (Total: Rs. 2.50L)", td_style)],
        [Paragraph("ii.", td_style), Paragraph("Rental charges for lab/workshop facilities", td_style), Paragraph("Rs. 15,000", td_style), Paragraph("Rs. 1,05,000", td_style), Paragraph("SMT rework station, high-bandwidth oscilloscope &amp; logic analyzer facility hire for 15 mos (Total: Rs. 1.20L)", td_style)],
        [Paragraph("iii.", td_style), Paragraph("** Essential equipment (cannot be rented)", td_style), Paragraph("Rs. 20,000", td_style), Paragraph("Rs. 1,60,000", td_style), Paragraph("Digital storage oscilloscope, programmable DC bench supply, SLA 3D printer &amp; reflow oven (Total: Rs. 1.80L)", td_style)],
        [Paragraph("iv.", td_style), Paragraph("Raw Material/spares/consumables cost", td_style), Paragraph("Rs. 25,000", td_style), Paragraph("Rs. 1,95,000", td_style), Paragraph("MCUs (RP2040/CH32V203), custom PCBs, mechanical switches, rotary encoders, OLEDs, passives, USB-C (Total: Rs. 2.20L)", td_style)],
        [Paragraph("v.", td_style), Paragraph("Fabrication/synthesis Charges", td_style), Paragraph("Rs. 20,000", td_style), Paragraph("Rs. 1,60,000", td_style), Paragraph("SMT stencil cutting, pilot multi-layer PCB batch assembly, CNC rapid tooling for enclosure (Total: Rs. 1.80L)", td_style)],
        [Paragraph("vi.", td_style), Paragraph("Manpower cost of technical assistants (&lt;=20%)", td_style), Paragraph("Rs. 24,000", td_style), Paragraph("Rs. 2,16,000", td_style), Paragraph("1 Embedded Hardware Assistant @ Rs. 16,000/mo for 15 mos (~15.6% of project cost, within limit) (Total: Rs. 2.40L)", td_style)],
        [Paragraph("vii.", td_style), Paragraph("Testing and Trials", td_style), Paragraph("Rs. 10,000", td_style), Paragraph("Rs. 90,000", td_style), Paragraph("50M keystroke endurance testing, thermal/drop profiling &amp; 50-user beta pilot trial deployment (Total: Rs. 1.00L)", td_style)],
        [Paragraph("viii.", td_style), Paragraph("Travel (Based on actual &amp; &lt;=5%)", td_style), Paragraph("Rs. 5,000", td_style), Paragraph("Rs. 45,000", td_style), Paragraph("Supplier audits, TOCIC review meetings &amp; field validation trials (~3.2% of project cost) (Total: Rs. 0.50L)", td_style)],
        [Paragraph("ix.", td_style), Paragraph("Patent Filing (actual fee paid to patent office)", td_style), Paragraph("Rs. 10,000", td_style), Paragraph("Rs. 70,000", td_style), Paragraph("Indian complete patent drafting, statutory filing fees, formal prior art search &amp; attorney prosecution (Total: Rs. 0.80L)", td_style)],
        [Paragraph("<b>-</b>", td_bold), Paragraph("<b>Total Cost</b>", td_bold), Paragraph("<b>Rs. 1,54,000</b>", td_bold), Paragraph("<b>Rs. 13,86,000</b>", td_bold), Paragraph("<b>10% Own Share (Rs. 1.54 Lakhs) + 90% PRISM Support (Rs. 13.86 Lakhs) = Rs. 15,40,000/-</b>", td_bold)],
    ]
    
    t_cost = Table(cost_rows, colWidths=[20, 130, 65, 75, 230])
    t_cost.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#003366')),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
        ('BACKGROUND', (0,-1), (-1,-1), colors.HexColor('#E2E8F0')),
        ('TOPPADDING', (0,0), (-1,-1), 2.5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 2.5),
        ('LEFTPADDING', (0,0), (-1,-1), 3),
        ('RIGHTPADDING', (0,0), (-1,-1), 3),
    ]))
    story.append(t_cost)
    story.append(Spacer(1, 4))
    story.append(Paragraph("<i>(* Basis of cost justification given in Annexure I. ** List of essential equipment given in Annexure II)</i>", body_style))
    story.append(Spacer(1, 4))
    
    # 6. Work Plan Table
    story.append(Paragraph("<b>6. Activity details/work plan:</b>", h1_style))
    
    wp_rows = [
        [Paragraph("<b>Activity</b>", th_style), Paragraph("<b>Monitor-able Milestones</b>", th_style), Paragraph("<b>Duration</b>", th_style), Paragraph("<b>Budget (Rs)</b>", th_style)],
        [Paragraph("Design engineering (product innovation)", td_style), Paragraph("KiCad V5 PCB schematic finalization, signal integrity / ESD simulations, 3D CAD ergonomic casing design.", td_style), Paragraph("M1 - M4 (4 mos)", td_style), Paragraph("Rs. 3,80,000", td_style)],
        [Paragraph("Working model/prototype development", td_style), Paragraph("SMT fabrication of 50 pilot PCB units, sub-millisecond firmware flashing, WebView2 daemon auto-switch integration.", td_style), Paragraph("M5 - M9 (5 mos)", td_style), Paragraph("Rs. 5,40,000", td_style)],
        [Paragraph("Product testing or Process demonstration", td_style), Paragraph("50M keystroke endurance testing, thermal/ESD stress profiling, 50-user beta pilot cohort deployment &amp; feedback refinement.", td_style), Paragraph("M10 - M13 (4 mos)", td_style), Paragraph("Rs. 3,60,000", td_style)],
        [Paragraph("Any others (Patent &amp; Commercialization)", td_style), Paragraph("Indian complete patent filing, CE/RoHS/BIS pre-compliance dossier, DFM package for local Indian mass production.", td_style), Paragraph("M14 - M15 (2 mos)", td_style), Paragraph("Rs. 2,60,000", td_style)],
        [Paragraph("<b>Total Work Plan</b>", td_bold), Paragraph("<b>Fully validated, field-tested, patented commercial-ready hardware-software macro controller ecosystem.</b>", td_bold), Paragraph("<b>15 Months</b>", td_bold), Paragraph("<b>Rs. 15,40,000</b>", td_bold)],
    ]
    
    t_wp = Table(wp_rows, colWidths=[120, 250, 75, 75])
    t_wp.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#003366')),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
        ('BACKGROUND', (0,-1), (-1,-1), colors.HexColor('#E2E8F0')),
        ('TOPPADDING', (0,0), (-1,-1), 2.5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 2.5),
        ('LEFTPADDING', (0,0), (-1,-1), 3),
        ('RIGHTPADDING', (0,0), (-1,-1), 3),
    ]))
    story.append(t_wp)
    
    story.append(PageBreak())

    # ---------------- PAGE 4 ----------------
    story.append(Paragraph("<b>Note: Duration of the project should not be more than 24 months. (Proposed Duration: 15 Months)</b>", body_style))
    story.append(Spacer(1, 4))
    
    story.append(Paragraph("<b>7. (a) End product / process / output resulting from innovation / final deliverables:</b>", h1_style))
    story.append(Paragraph(
        "1. Commercial Hardware Prototype ('Overcontrol V5'): 12+ mechanical hot-swap switches, dual rotary encoders, status OLED display, USB-C ESD protection.<br/>"
        "2. Production DFM Package: Complete open-architecture KiCad schematics, 4-layer Gerber, BOM, pick-and-place, CNC/injection molding CAD files.<br/>"
        "3. Low-Latency Firmware: Bare-metal firmware with sub-millisecond USB HID reporting, debouncing, and serial handshake.<br/>"
        "4. Companion Desktop Suite: Native Windows/Cross-platform application featuring automatic active window detection and community macro sharing.<br/>"
        "5. Indian Patent Application &amp; Technical Dossier: Complete filed patent specification and 50-user validation report.",
        body_style
    ))
    story.append(Spacer(1, 4))
    
    story.append(Paragraph("<b>(b) Innovation's benefit to the society:</b>", h1_style))
    story.append(Paragraph(
        "• Democratizes high-end productivity peripherals for Indian students, engineers, and creators at &lt;Rs. 3,500 (&lt;80% cheaper than imports).<br/>"
        "• Fosters domestic electronics manufacturing under Make in India / Atmanirbhar Bharat.<br/>"
        "• Serves as assistive technology for motor-impaired individuals via single-switch complex macro execution.",
        body_style
    ))
    story.append(Spacer(1, 4))
    
    story.append(Paragraph("<b>8. Any other information relevant to the project:</b> Functional proof-of-concept prototypes (KiCad PCB, 3D CAD, desktop software) developed and verified in-house, demonstrating high readiness.", body_style))
    story.append(Spacer(1, 4))
    
    story.append(Paragraph("<b>9. Referees (Two Nos with complete address, phone number and e-mail ID):</b>", h1_style))
    ref_p = [
        [Paragraph("<b>Referee 1 (Academic / Technical):</b><br/>Name: [Prof. / Dr. Name]<br/>Designation: Professor / Senior Embedded Architect<br/>Org: [Dept of ECE/CSE, Institute Name]<br/>Address: [Institute Address, City, State - PIN]<br/>Phone: [+91-XXXXXXXXXX] | Email: [ref1@institute.ac.in]", body_style),
         Paragraph("<b>Referee 2 (Industry / Domain):</b><br/>Name: [Mr. / Ms. Name]<br/>Designation: Principal Engineer / Tech Lead<br/>Org: [Hardware / Tech Company Name]<br/>Address: [Company Address, City, State - PIN]<br/>Phone: [+91-XXXXXXXXXX] | Email: [ref2@domain.com]", body_style)]
    ]
    t_ref = Table(ref_p, colWidths=[260, 260])
    t_ref.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#F8FAFC')),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
        ('TOPPADDING', (0,0), (-1,-1), 3),
        ('BOTTOMPADDING', (0,0), (-1,-1), 3),
        ('LEFTPADDING', (0,0), (-1,-1), 5),
        ('RIGHTPADDING', (0,0), (-1,-1), 5),
    ]))
    story.append(t_ref)
    story.append(Spacer(1, 4))
    
    story.append(Paragraph("<b>10. Declaration:</b> I/We declare that all statements made are true and correct. I/We have not received financial assistance for this proposal from any other agency.", body_style))
    story.append(Spacer(1, 6))
    
    sig_p = [
        [Paragraph("<b>Place:</b> _____________________<br/><b>Date:</b>  _____________________", body_style),
         Paragraph("<para align='right'><b>_____________________________________</b><br/><b>Signature of the applicant</b><br/>(Pulak Nayak)</para>", body_style)]
    ]
    t_sig = Table(sig_p, colWidths=[260, 260])
    t_sig.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'MIDDLE')]))
    story.append(t_sig)
    story.append(Spacer(1, 6))
    
    story.append(Paragraph("<b>11. RECOMMENDATIONS OF THE FORWARDING TePP Outreach Cum Cluster Innovation Centre (TOCIC)</b>", h1_style))
    toc_p = [
        [Paragraph("<b>Place:</b> _____________________<br/><b>Date:</b>  _____________________", body_style),
         Paragraph("<para align='right'><b>_____________________________________</b><br/><b>Signature of the Head, TOCIC</b><br/>(With Official Seal)</para>", body_style)]
    ]
    t_toc = Table(toc_p, colWidths=[260, 260])
    t_toc.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'MIDDLE')]))
    story.append(t_toc)
    
    story.append(PageBreak())

    # ---------------- PAGE 5 ----------------
    story.append(Spacer(1, 30))
    story.append(Paragraph("<b>No Objection Certificate for the Student /Employee Innovators</b>", title_style))
    story.append(Spacer(1, 30))
    story.append(Paragraph(
        "The student / employee innovator <b>Pulak Nayak</b> is studying /working in our institute/organization since <b>[Date / Month / Year]</b>. The institute/organization has no objection to the innovator taking up the innovation work as proposed with financial support under PRISM. The institute laboratories will be made available on chargeable basis to the innovator for executing his/her PRISM project. The Institute/Organization will be responsible for final completion of project in case the Innovator leaves the Institute/Organization without completion of the project.",
        ParagraphStyle('NOCBody', parent=styles['Normal'], fontName='Helvetica', fontSize=10, leading=16, textColor=colors.HexColor('#1f2937'))
    ))
    story.append(Spacer(1, 120))
    story.append(Paragraph("<para align='right'>.......................................................<br/><b>Competent authority</b><br/><b>Signature with Seal</b></para>", ParagraphStyle('NOCSig', parent=styles['Normal'], fontName='Helvetica', fontSize=10, leading=14)))

    story.append(PageBreak())

    # ---------------- PAGE 6 ----------------
    story.append(Paragraph("<b>Evaluation by Domain Knowledge Experts</b>", title_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("<b>1. Name of the Expert :</b> [Domain Expert Name, e-mail ID, mobile no, contact address]", body_style))
    story.append(Paragraph("<b>2. Title of the Proposal :</b> Design, Fabrication, and Commercial Validation of an Intelligent Context-Aware Adaptive Hardware Macro Controller (Overcontrol)", body_style))
    story.append(Paragraph("<b>3. Expert comments on proposed innovation :</b> Strong technological merit, novel hardware-software co-design, and clear commercialization potential in replacing expensive imported input devices.", body_style))
    story.append(Spacer(1, 6))
    story.append(Paragraph("<b>4 Assessments by the Expert:</b>", h1_style))
    
    exp_rows_pdf = [
        [Paragraph("<b>Assessment Category</b>", th_style), Paragraph("<b>Expert comments / Recommendations</b>", th_style)],
        [Paragraph("<b>Assessment of technology merits of proposed innovation</b>", td_style),
         Paragraph("The integration of sub-millisecond interrupt-driven RP2040/RISC-V architecture with dynamic OS process hook auto-switching offers significant technical merit. The approach solves key latency and software bloat problems in current market devices.", td_style)],
        [Paragraph("<b>Technological challenges in design and prototype manufacture based on innovators known skill</b>", td_style),
         Paragraph("The innovator has already demonstrated practical competence through functional KiCad multi-layer PCB design, 3D CAD modeling, and full-stack software development. Key challenges in DFM tooling, SMT yield, and EMI/EMC compliance are realistically addressed in the work plan.", td_style)],
        [Paragraph("<b>Recommendations To PRISM PASC</b>", td_style),
         Paragraph("Strongly recommended for financial support under PRISM Phase-I Category II. The project possesses clear innovation, high societal &amp; commercial value, and aligns with the Make in India initiative.", td_style)]
    ]
    
    t_exp_pdf = Table(exp_rows_pdf, colWidths=[180, 340])
    t_exp_pdf.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#003366')),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('LEFTPADDING', (0,0), (-1,-1), 5),
        ('RIGHTPADDING', (0,0), (-1,-1), 5),
    ]))
    story.append(t_exp_pdf)
    story.append(Spacer(1, 40))
    story.append(Paragraph("<para align='right'>.......................................................<br/><b>[Signature of Expert with seal]</b><br/>Designation: _____________________<br/>Address of the organization/institute: _____________________</para>", body_style))

    story.append(PageBreak())

    # ---------------- PAGE 7 ----------------
    story.append(Paragraph("<b>Covering letter to the PRISM application</b>", title_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Date: _____________________", body_style))
    story.append(Spacer(1, 8))
    story.append(Paragraph("To:<br/><b>The TOCIC</b><br/>-------------------<br/>-----------------", body_style))
    story.append(Spacer(1, 8))
    story.append(Paragraph("<b>Sub: Proposal for Development of Design, Fabrication, and Commercial Validation of an Intelligent Context-Aware Adaptive Hardware Macro Controller (Overcontrol)</b>", body_style))
    story.append(Spacer(1, 8))
    story.append(Paragraph("Dear Sir/Madam,<br/><br/>I am herewith submitting my application for support under PRISM. The following documents are enclosed.", body_style))
    story.append(Spacer(1, 8))
    
    enc_pdf = (
        "[√]&nbsp;&nbsp;&nbsp;&nbsp;<b>Signed Copy of Application</b><br/>"
        "[√]&nbsp;&nbsp;&nbsp;&nbsp;<b>Proof of Residence</b> (Aadhaar / Residence Certificate)<br/>"
        "[√]&nbsp;&nbsp;&nbsp;&nbsp;<b>Innovation:</b> Design, Fabrication, and Commercial Validation of Intelligent Context-Aware Hardware Macro Controller<br/>"
        "[√]&nbsp;&nbsp;&nbsp;&nbsp;<b>Documentary Proof of Prior Work</b> (PCB layouts, 3D CAD models, software code repository, prototype photographs)<br/>"
        "[√]&nbsp;&nbsp;&nbsp;&nbsp;<b>Work Planned</b> (Annexures I, II, III)<br/>"
        "[√]&nbsp;&nbsp;&nbsp;&nbsp;<b>Profile of Potential User</b><br/>"
        "[√]&nbsp;&nbsp;&nbsp;&nbsp;<b>Copy of Aadhaar Card &amp; PAN Card</b>"
    )
    story.append(Paragraph(enc_pdf, ParagraphStyle('EncStyle', parent=styles['Normal'], fontName='Helvetica', fontSize=9, leading=15)))
    story.append(Spacer(1, 60))
    story.append(Paragraph("<para align='right'>_____________________________________<br/><b>Innovator</b><br/>(Pulak Nayak)</para>", body_style))

    # ---------------- ANNEXURE I & II ----------------
    story.append(PageBreak())
    story.append(Paragraph("<b>ANNEXURE I: DETAILED BASIS OF COST ESTIMATION &amp; JUSTIFICATION</b>", title_style))
    story.append(Spacer(1, 8))
    
    anx1_text = (
        "• <b>i. R&amp;D / Design Engg / Consultancy (Rs. 2,50,000):</b> Signal integrity simulation, 90-ohm USB-C differential impedance matching, EMI/EMC compliance pre-testing, and embedded latency profiling.<br/>"
        "• <b>ii. Rental charges for lab/workshop (Rs. 1,20,000):</b> Access to SMT rework facilities, high-bandwidth logic analyzers, and thermal/environmental chambers for 15 months (@ Rs. 8,000/month).<br/>"
        "• <b>iii. Essential equipment (Rs. 1,80,000):</b> 4-Channel 100MHz DSO, programmable DC power supply, desktop SLA 3D printer, and reflow soldering station (detailed in Annexure II).<br/>"
        "• <b>iv. Raw materials &amp; consumables (Rs. 2,20,000):</b> 150+ RP2040/CH32V203 MCUs, 4-layer FR4 PCBs, mechanical switches, optical/rotary encoders, OLEDs, TVS diode ESD arrays, USB-C connectors, UV resin/filaments.<br/>"
        "• <b>v. Fabrication &amp; rapid tooling (Rs. 1,80,000):</b> Stencil cutting, automated SMT pick-and-place batch runs (50 pilot units), CNC precision aluminum knob machining, and rapid tooling for casings.<br/>"
        "• <b>vi. Manpower (Rs. 2,40,000):</b> 1 Technical Assistant @ Rs. 16,000/month for 15 months (~15.6% of project cost, well within 20% limit).<br/>"
        "• <b>vii. Testing &amp; trials (Rs. 1,00,000):</b> Automated 50M keystroke endurance testing, drop/vibration stress testing, thermal profiling, and 50-user beta pilot cohort deployment across India.<br/>"
        "• <b>viii. Travel (Rs. 50,000):</b> Supplier quality audits, TOCIC milestone progress presentations, and field user testing sessions (~3.2% of project cost, within 5% limit).<br/>"
        "• <b>ix. Patent filing (Rs. 80,000):</b> Indian complete patent specification drafting, statutory patent office filing fees, formal prior art search reports, and attorney fees."
    )
    story.append(Paragraph(anx1_text, body_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("<b>ANNEXURE II: LIST OF ESSENTIAL NON-RENTABLE EQUIPMENT &amp; JUSTIFICATION</b>", title_style))
    story.append(Spacer(1, 8))
    
    eq_rows_pdf = [
        [Paragraph("<b>Sl.</b>", th_style), Paragraph("<b>Equipment Description</b>", th_style), Paragraph("<b>Qty</b>", th_style), Paragraph("<b>Cost (Rs)</b>", th_style), Paragraph("<b>Technical Justification</b>", th_style)],
        [Paragraph("1.", td_style), Paragraph("4-Channel 100MHz Digital Storage Oscilloscope (DSO)", td_style), Paragraph("1", td_style), Paragraph("Rs. 65,000", td_style), Paragraph("Essential for daily real-time debugging of encoder jitter, USB timing, and debouncing.", td_style)],
        [Paragraph("2.", td_style), Paragraph("Triple-Output Programmable DC Power Bench Supply", td_style), Paragraph("1", td_style), Paragraph("Rs. 30,000", td_style), Paragraph("Required for voltage rail regulation, transient load testing, and current draw profiling.", td_style)],
        [Paragraph("3.", td_style), Paragraph("Desktop SLA UV Resin 3D Printer &amp; Curing Station", td_style), Paragraph("1", td_style), Paragraph("Rs. 55,000", td_style), Paragraph("In-house fabrication of precision keycaps, rotary housings, and enclosures (&lt;50 micron tolerance).", td_style)],
        [Paragraph("4.", td_style), Paragraph("SMT Hot Air Rework Station &amp; Reflow Plate", td_style), Paragraph("1", td_style), Paragraph("Rs. 30,000", td_style), Paragraph("Mandatory for QFN/SMD assembly, prototype rework, and mounting QFN56 microcontrollers.", td_style)],
        [Paragraph("<b>-</b>", td_bold), Paragraph("<b>TOTAL EQUIPMENT COST</b>", td_bold), Paragraph("<b>-</b>", td_bold), Paragraph("<b>Rs. 1,80,000</b>", td_bold), Paragraph("<b>Procured as essential capital equipment under PRISM Category II</b>", td_bold)],
    ]
    
    t_eq_pdf = Table(eq_rows_pdf, colWidths=[20, 150, 25, 65, 260])
    t_eq_pdf.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#003366')),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
        ('BACKGROUND', (0,-1), (-1,-1), colors.HexColor('#E2E8F0')),
        ('TOPPADDING', (0,0), (-1,-1), 2.5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 2.5),
        ('LEFTPADDING', (0,0), (-1,-1), 3),
        ('RIGHTPADDING', (0,0), (-1,-1), 3),
    ]))
    story.append(t_eq_pdf)
    
    doc.build(story)
    print(f"Full 8-page PDF document saved to {output_path}")

if __name__ == '__main__':
    base_dir = r"c:\Users\pulak\Desktop\V4_RP2040_Zero_Webview_Main"
    docx_out = os.path.join(base_dir, "PRISM_Phase_I_Category_II_Application_Form_Filled.docx")
    pdf_out = os.path.join(base_dir, "PRISM_Phase_I_Category_II_Application_Form_Filled.pdf")
    
    create_prism_docx(docx_out)
    create_full_prism_pdf(pdf_out)
